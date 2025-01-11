import os
import json
from openai import OpenAI
import PyPDF2
from docx import Document as DocxDocument
import io
from typing import Dict, Any, List, Union, Optional
from openai.types.chat import ChatCompletionMessageParam, ChatCompletionSystemMessageParam, ChatCompletionUserMessageParam
from PyPDF2 import PdfReader
from docx import Document
import openai
from models import db, FEMAForm, FEMARequirement, FEMAAnalysis

# Initialize OpenAI client
client = OpenAI()

# Load prompts from JSON file
def load_prompts() -> Dict[str, Any]:
    with open('prompts.json', 'r') as f:
        return json.load(f)

# Global prompts dictionary
PROMPTS = load_prompts()

def process_document(file, processing_type='text') -> Dict[str, Any]:
    """Process document with specified processing type."""
    filename = file.filename.lower()
    
    # Extract raw text content
    if filename.endswith('.pdf'):
        content = process_pdf(file)
    elif filename.endswith('.docx'):
        content = process_docx(file)
    elif filename.endswith('.txt'):
        content = process_txt(file)
    else:
        raise ValueError("Unsupported file format")
    
    # Process content based on type
    if processing_type == 'text':
        return {'raw_text': content, 'processed_text': content}
    elif processing_type == 'summary':
        return process_summary(content)
    elif processing_type == 'analysis':
        return process_analysis(content)
    elif processing_type == 'insurance_requirements':
        return process_insurance_requirements(content)
    else:
        raise ValueError("Unsupported processing type")

def process_pdf(file) -> str:
    """Extract text from PDF files."""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error processing PDF: {str(e)}")

def process_docx(file) -> str:
    """Extract text from DOCX files."""
    try:
        doc = DocxDocument(io.BytesIO(file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error processing DOCX: {str(e)}")

def process_txt(file) -> str:
    """Extract text from TXT files."""
    try:
        return file.read().decode('utf-8').strip()
    except Exception as e:
        raise ValueError(f"Error processing TXT: {str(e)}")

def process_summary(content: str) -> Dict[str, Any]:
    """Generate a summary of the document using OpenAI."""
    try:
        messages: List[ChatCompletionMessageParam] = [
            ChatCompletionSystemMessageParam(
                role=PROMPTS["document_summary"]["role"],
                content=PROMPTS["document_summary"]["content"]
            ),
            ChatCompletionUserMessageParam(
                role="user",
                content=content
            )
        ]

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            response_format={"type": "json_object"}
        )
        
        if response.choices and response.choices[0].message and response.choices[0].message.content:
            return json.loads(response.choices[0].message.content)
        raise Exception("Empty response from OpenAI")
    except Exception as e:
        raise Exception(f"Failed to generate summary: {str(e)}")

def process_analysis(content: str) -> Dict[str, Any]:
    """Analyze the document content using OpenAI."""
    try:
        messages: List[ChatCompletionMessageParam] = [
            ChatCompletionSystemMessageParam(
                role=PROMPTS["document_analysis"]["role"],
                content=PROMPTS["document_analysis"]["content"]
            ),
            ChatCompletionUserMessageParam(
                role="user",
                content=content
            )
        ]

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            response_format={"type": "json_object"}
        )
        
        if response.choices and response.choices[0].message and response.choices[0].message.content:
            return json.loads(response.choices[0].message.content)
        raise Exception("Empty response from OpenAI")
    except Exception as e:
        raise Exception(f"Failed to analyze content: {str(e)}")

def process_insurance_requirements(content: str) -> Dict[str, Any]:
    """Extract and analyze insurance requirements using OpenAI."""
    try:
        messages: List[ChatCompletionMessageParam] = [
            ChatCompletionSystemMessageParam(
                role=PROMPTS["insurance_requirements"]["role"],
                content=PROMPTS["insurance_requirements"]["content"]
            ),
            ChatCompletionUserMessageParam(
                role="user",
                content=content
            )
        ]

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            response_format={"type": "json_object"}
        )
        
        if response.choices and response.choices[0].message and response.choices[0].message.content:
            return json.loads(response.choices[0].message.content)
        raise Exception("Empty response from OpenAI")
    except Exception as e:
        raise Exception(f"Failed to process insurance requirements: {str(e)}")

def analyze_insurance_claim(claim_content: str, requirement_content: str, analysis_type: str, previous_messages: List[Dict[str, str]] = None) -> Dict[str, Any]:
    """Analyze insurance claim against requirements with specified analysis type."""
    try:
        # Map frontend analysis types to backend types
        analysis_type_map = {
            'explain': 'explain',
            'enhance': 'enhance',
            'mock_rejection': 'mock_rejection',
            'language': 'language'
        }
        
        backend_analysis_type = analysis_type_map.get(analysis_type, analysis_type)
        
        if backend_analysis_type not in PROMPTS["insurance_analysis"]:
            raise ValueError(f"Invalid analysis type: {backend_analysis_type}")

        # Build context from previous messages if they exist
        context_prompt = ""
        if previous_messages:
            context_prompt = "\nPrevious analyses have already covered these points:\n"
            for msg in previous_messages:
                if msg["role"] == "assistant" and msg.get("content"):
                    try:
                        content = json.loads(msg["content"])
                        if "analysis" in content and "improvements" in content["analysis"]:
                            for priority, items in content["analysis"]["improvements"].items():
                                context_prompt += f"\n{priority}:\n"
                                context_prompt += "\n".join(items) + "\n"
                    except json.JSONDecodeError:
                        continue
            context_prompt += "\nProvide 25 NEW and UNIQUE points that haven't been covered above.\n"

        messages: List[ChatCompletionMessageParam] = [
            ChatCompletionSystemMessageParam(
                role=PROMPTS["insurance_analysis"][backend_analysis_type]["role"],
                content=f"""{PROMPTS["insurance_analysis"][backend_analysis_type]["content"]}
{context_prompt}
Provide your response as exactly 25 NEW items, organized into priority sections but maintaining a sequential numbering from 1-25. Format as follows:

{{
    "analysis": {{
        "summary": "detailed analysis text",
        "score": 0-100,
        "improvements": {{
            "high_priority": [
                "1. First high priority item",
                "2. Second high priority item",
                "3. Third high priority item"
            ],
            "medium_priority": [
                "8. First medium priority item",
                "9. Second medium priority item"
            ],
            "low_priority": [
                "15. First low priority item",
                "16. Second low priority item"
            ]
        }}
    }},
    "details": {{
        "matching_requirements": ["req1", "req2"],
        "missing_requirements": ["req1", "req2"],
        "compliance_score": 0-100,
        "is_followup_analysis": true
    }}
}}

IMPORTANT: 
- Maintain sequential numbering from 1 to 25 across all sections
- Each item should start with its number (1-25) followed by a period
- The total number of items across all priority sections must equal exactly 25
- All items must be NEW and not mentioned in any previous analyses
- Focus on different aspects or deeper details than previous analyses
- Ensure your response is a valid JSON object"""
            ),
            ChatCompletionUserMessageParam(
                role="user",
                content=f"Requirements:\n{requirement_content}\n\nClaim:\n{claim_content}"
            )
        ]

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            response_format={"type": "json_object"}
        )
        
        if response.choices and response.choices[0].message and response.choices[0].message.content:
            raw_content = response.choices[0].message.content
            print(f"Raw AI response: {raw_content}")  # Log the raw response
            try:
                return json.loads(raw_content)
            except json.JSONDecodeError as json_error:
                print(f"JSON parsing error: {str(json_error)}")
                print(f"Problematic content: {raw_content}")
                # Attempt to clean and parse the response
                cleaned_content = raw_content.strip().replace('\n', '').replace('\r', '')
                try:
                    return json.loads(cleaned_content)
                except json.JSONDecodeError:
                    raise ValueError(f"Failed to parse JSON response even after cleaning. Raw content: {raw_content}")
        raise Exception("Empty response from OpenAI")
    except Exception as e:
        print(f"Error in analyze_insurance_claim: {str(e)}")
        raise Exception(f"Failed to analyze insurance claim: {str(e)}")

def process_fema_document(file_path: Optional[str], analysis_type: str, form_id: Optional[int] = None) -> Union[List[str], Dict]:
    """Process FEMA documents for various analysis types."""
    print(f"Starting document processing - Type: {analysis_type}, Path: {file_path}")  # Debug log
    
    # Load prompts
    try:
        with open('prompts.json', 'r') as f:
            prompts = json.load(f)
    except Exception as e:
        print(f"Error loading prompts: {str(e)}")  # Debug log
        raise Exception("Failed to load processing prompts")

    if analysis_type == 'requirements':
        if not file_path or not os.path.exists(file_path):
            print(f"File not found: {file_path}")  # Debug log
            raise ValueError("File not found or invalid file path")

        # Extract requirements from uploaded document
        try:
            text = extract_text_from_document(file_path)
            print(f"Extracted text length: {len(text)}")  # Debug log
            requirements = analyze_requirements(text, prompts['fema_requirements'])
            print(f"Analyzed requirements count: {len(requirements)}")  # Debug log
            return requirements
        except Exception as e:
            print(f"Error in requirements processing: {str(e)}")  # Debug log
            raise Exception(f"Failed to process requirements: {str(e)}")

    # For other analysis types, we need the form_id
    if not form_id:
        raise ValueError("form_id is required for this analysis type")

    # Get the form and its requirements
    form = FEMAForm.query.get(form_id)
    if not form:
        raise ValueError("Invalid form ID")

    requirements = [req.requirement_text for req in form.requirements]
    
    # Get the form text if we're analyzing a new upload
    form_text = ""
    if file_path:
        try:
            form_text = extract_text_from_document(file_path)
        except Exception as e:
            print(f"Error extracting text from form: {str(e)}")  # Debug log
            raise Exception(f"Failed to extract text from form: {str(e)}")
    else:
        # Get the latest analysis of type 'form' for this form_id
        form_analysis = FEMAAnalysis.query.filter_by(
            fema_form_id=form_id,
            analysis_type='form'
        ).order_by(FEMAAnalysis.created_at.desc()).first()
        
        if form_analysis:
            form_text = form_analysis.content.get('text', '')

    if not form_text:
        raise ValueError("No form text available for analysis")

    # Perform the requested analysis
    try:
        if analysis_type == 'form':
            return {
                'text': form_text,
                'analysis': analyze_form(form_text, requirements, prompts['fema_form_analysis'])
            }
        elif analysis_type == 'explanation':
            return analyze_form_explanation(form_text, prompts['fema_form_explanation'])
        elif analysis_type == 'enhancement':
            return analyze_form_enhancement(form_text, requirements, prompts['fema_form_enhancement'])
        elif analysis_type == 'rejection':
            return analyze_rejection_reasons(form_text, requirements, prompts['fema_form_rejection'])
        elif analysis_type == 'language':
            return analyze_language(form_text, prompts['fema_form_language'])
        else:
            raise ValueError(f"Unknown analysis type: {analysis_type}")
    except Exception as e:
        print(f"Error in analysis: {str(e)}")  # Debug log
        raise Exception(f"Failed to analyze document: {str(e)}")

def extract_text_from_document(file_path: str) -> str:
    """Extract text from PDF, DOCX, or TXT files."""
    print(f"Extracting text from: {file_path}")  # Debug log
    
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    try:
        if ext == '.pdf':
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        elif ext == '.docx':
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif ext == '.txt':
            with open(file_path, 'r') as f:
                text = f.read()
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        print(f"Successfully extracted {len(text)} characters")  # Debug log
        return text
    except Exception as e:
        print(f"Error extracting text: {str(e)}")  # Debug log
        raise Exception(f"Failed to extract text from file: {str(e)}")

def analyze_requirements(text: str, prompt: str) -> List[str]:
    """Extract requirements from document text."""
    print("Starting requirements analysis")  # Debug log
    
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": prompt},
                {"role": "user", "content": text}
            ],
            temperature=0.7,
            max_tokens=2000
        )

        # Extract requirements from the response
        requirements = []
        for line in response.choices[0].message.content.split('\n'):
            line = line.strip()
            if line and not line.startswith('#') and not line.startswith('-'):
                requirements.append(line)

        print(f"Extracted {len(requirements)} requirements")  # Debug log
        return requirements
    except Exception as e:
        print(f"Error in requirements analysis: {str(e)}")  # Debug log
        raise Exception(f"Failed to analyze requirements: {str(e)}")

def analyze_form(text: str, requirements: List[str], prompt: str) -> Dict:
    """Analyze the form against requirements."""
    
    context = f"Requirements:\n" + "\n".join(f"- {req}" for req in requirements)
    context += f"\n\nForm Text:\n{text}"

    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": context}
        ],
        temperature=0.7,
        max_tokens=2000
    )

    return json.loads(response.choices[0].message.content)

def analyze_form_explanation(text: str, prompt: str) -> Dict:
    """Generate a detailed explanation of the form content."""
    
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": text}
        ],
        temperature=0.7,
        max_tokens=2000
    )

    return json.loads(response.choices[0].message.content)

def analyze_form_enhancement(text: str, requirements: List[str], prompt: str) -> Dict:
    """Suggest enhancements for the form."""
    
    context = f"Requirements:\n" + "\n".join(f"- {req}" for req in requirements)
    context += f"\n\nForm Text:\n{text}"

    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": context}
        ],
        temperature=0.7,
        max_tokens=2000
    )

    return json.loads(response.choices[0].message.content)

def analyze_rejection_reasons(text: str, requirements: List[str], prompt: str) -> Dict:
    """Identify potential reasons for form rejection."""
    
    context = f"Requirements:\n" + "\n".join(f"- {req}" for req in requirements)
    context += f"\n\nForm Text:\n{text}"

    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": context}
        ],
        temperature=0.7,
        max_tokens=2000
    )

    return json.loads(response.choices[0].message.content)

def analyze_language(text: str, prompt: str) -> Dict:
    """Analyze the language and writing style of the form."""
    
    response = openai.ChatCompletion.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": text}
        ],
        temperature=0.7,
        max_tokens=2000
    )

    return json.loads(response.choices[0].message.content)
